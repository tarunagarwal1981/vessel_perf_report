import os
import streamlit as st
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ReportIconManager:
    def __init__(self, icons_folder="icons"):
        """
        Initialize with the folder containing report icons
        
        Parameters:
        icons_folder (str): Path to the folder containing the icon images
        """
        self.icons_folder = icons_folder
        
        # Define the standard icon paths
        self.icon_paths = {
            'hull': os.path.join(icons_folder, "hull_icon.png"),
            'machinery': os.path.join(icons_folder, "machinery_icon.png"),
            'emissions': os.path.join(icons_folder, "emissions_icon.png"),
            'green_circle': os.path.join(icons_folder, "green_circle.png"),
            'orange_circle': os.path.join(icons_folder, "orange_circle.png"),
            'red_circle': os.path.join(icons_folder, "red_circle.png")
        }
    
    def add_icon_to_cell(self, cell, icon_key, width_inches=0.7):
        """
        Add an icon to a table cell
        
        Parameters:
        cell: The document cell to add the icon to
        icon_key (str): Key to identify which icon to use (from icon_paths)
        width_inches (float): Width of the icon in inches
        
        Returns:
        bool: True if successful, False if failed
        """
        try:
            if icon_key in self.icon_paths and os.path.exists(self.icon_paths[icon_key]):
                paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(self.icon_paths[icon_key], width=Inches(width_inches))
                return True
            else:
                # Add placeholder text if icon not found
                if not cell.paragraphs:
                    cell.add_paragraph()
                cell.paragraphs[0].text = f"[{icon_key} icon]"
                return False
        except Exception as e:
            print(f"Error adding icon to cell: {str(e)}")
            cell.paragraphs[0].text = "[Icon Error]"
            return False
    
    def get_condition_icon(self, condition):
        """
        Get the appropriate icon key based on condition
        
        Parameters:
        condition (str): Condition string (GOOD, AVERAGE, POOR)
        
        Returns:
        str: Icon key to use
        """
        condition = condition.upper() if condition else ""
        if condition == "GOOD":
            return "green_circle"
        elif condition == "AVERAGE":
            return "orange_circle"
        elif condition == "POOR":
            return "red_circle"
        return None
    
    def add_condition_icon(self, cell, condition, width_inches=0.7):
        """
        Add the appropriate condition icon based on the condition value
        
        Parameters:
        cell: The document cell to add the icon to
        condition (str): Condition value (GOOD, AVERAGE, POOR)
        width_inches (float): Width of the icon in inches
        
        Returns:
        bool: True if successful, False if failed
        """
        icon_key = self.get_condition_icon(condition)
        if icon_key:
            return self.add_icon_to_cell(cell, icon_key, width_inches)
        else:
            # Add placeholder text if no appropriate icon
            if not cell.paragraphs:
                cell.add_paragraph()
            cell.paragraphs[0].text = f"[{condition} icon]"
            return False
