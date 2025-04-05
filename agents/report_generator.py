import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
import base64
import tempfile
import os
import copy
import uuid
import traceback

class ReportGeneratorAgent:
    def __init__(self):
        pass
    
    def run(self, vessel_data, selected_vessel, hull_agent, speed_agent):
        st.header("Performance Report Generator")
        
        # Check for valid data first
        if not vessel_data or len(vessel_data) == 0:
            st.warning("No data available for this vessel. Please select a different vessel or date range.")
            return
            
        # Create tabs for report view and download
        report_tab1, report_tab2 = st.tabs(["Report Preview", "Download Report"])
        
        with report_tab1:
            # Check if we have enough data
            if len(vessel_data) < 5:
                st.warning("Limited data available. The report may not be comprehensive.")
            
            try:
                # Get hull performance data
                hull_metrics = self._extract_hull_metrics(vessel_data, hull_agent)
                
                # Display report preview
                st.subheader("Vessel Performance Summary")
                
                # Create two columns for the summary data
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("### Hull & Propeller Performance")
                    st.markdown(f"**Additional Power Consumption:** {hull_metrics['power_loss']:.1f} %")
                    st.markdown(f"**Potential Fuel Savings:** {hull_metrics['fuel_savings']:.1f} MT/D")
                    st.markdown(f"**Hull & Propeller Condition:** {hull_metrics['condition']}")
                    st.markdown(f"**Recommendation:** {hull_metrics['recommendation']}")
                
                with col2:
                    st.markdown("### Speed Consumption Profile")
                    st.markdown("Average consumption rates at various speeds:")
                    
                    # Get speed-consumption data if available
                    speed_metrics = self._extract_speed_metrics(vessel_data)
                    if speed_metrics and 'ballast_data' in speed_metrics and speed_metrics['ballast_data']:
                        ballast_speeds = [round(x, 1) for x in speed_metrics['ballast_speeds']]
                        ballast_consumptions = [round(x, 1) for x in speed_metrics['ballast_consumptions']]
                        
                        # Display a small table
                        speed_df = pd.DataFrame({
                            "Speed (knots)": ballast_speeds[:5],  # Show just the first 5 for preview
                            "Consumption (MT/day)": ballast_consumptions[:5]
                        })
                        st.dataframe(speed_df, hide_index=True)
                
                # Show example charts
                st.subheader("Hull & Propeller Performance Analysis")
                if 'hull_chart' in hull_metrics and hull_metrics['hull_chart'] is not None:
                    # Generate a unique key for the chart
                    chart_key = f"hull_chart_{str(uuid.uuid4())[:8]}"
                    st.plotly_chart(hull_metrics['hull_chart'], use_container_width=True, key=chart_key)
                
                st.subheader("Speed Consumption Profile")
                # Create two columns for the speed consumption charts
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("**Ballast Condition**")
                    if 'ballast_chart' in speed_metrics and speed_metrics['ballast_chart'] is not None:
                        # Generate a unique key for the chart
                        ballast_key = f"ballast_chart_{str(uuid.uuid4())[:8]}"
                        st.plotly_chart(speed_metrics['ballast_chart'], use_container_width=True, key=ballast_key)
                    else:
                        st.info("No ballast condition data available")
                
                with col2:
                    st.markdown("**Laden Condition**")
                    if 'laden_chart' in speed_metrics and speed_metrics['laden_chart'] is not None:
                        # Generate a unique key for the chart
                        laden_key = f"laden_chart_{str(uuid.uuid4())[:8]}"
                        st.plotly_chart(speed_metrics['laden_chart'], use_container_width=True, key=laden_key)
                    else:
                        st.info("No laden condition data available")
            except Exception as e:
                st.error(f"Error in preview section: {str(e)}")
                st.code(traceback.format_exc())
        
        with report_tab2:
            st.markdown("### Generate and Download Report")
            
            # Input fields for report customization
            report_date = st.date_input("Report Date")
            analyst_name = st.text_input("Analyst Name", "Marine Performance Team")
            
            # Options for report content
            st.subheader("Report Content Options")
            include_hull = st.checkbox("Include Hull Performance Section", value=True)
            include_speed = st.checkbox("Include Speed-Consumption Section", value=True)
            include_emissions = st.checkbox("Include Emissions Section", value=False)
            include_machinery = st.checkbox("Include Machinery Section", value=False)
            
            # Option to include charts (can disable if causing issues)
            include_charts = st.checkbox("Include Charts in Report", value=True)
            
            # Warning about chart generation
            if include_charts:
                st.info("Note: Chart generation may fail in some environments. If you encounter issues, try disabling charts.")
            
            # Generate report button
            if st.button("Generate DOCX Report"):
                try:
                    with st.spinner("Generating report..."):
                        # Extract data for report
                        hull_metrics = self._extract_hull_metrics(vessel_data, hull_agent)
                        speed_metrics = self._extract_speed_metrics(vessel_data)
                        
                        # Create the DOCX report
                        docx_file = self._generate_docx_report(
                            vessel_name=selected_vessel,
                            report_date=report_date,
                            analyst_name=analyst_name,
                            hull_metrics=hull_metrics,
                            speed_metrics=speed_metrics,
                            options={
                                'include_hull': include_hull,
                                'include_speed': include_speed,
                                'include_emissions': include_emissions,
                                'include_machinery': include_machinery,
                                'include_charts': include_charts
                            }
                        )
                        
                        # Create download button
                        st.download_button(
                            label="Download DOCX Report",
                            data=docx_file,
                            file_name=f"{selected_vessel}_Performance_Report_{report_date}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        st.success("Report generated successfully!")
                except Exception as e:
                    st.error(f"Error generating report: {str(e)}")
                    st.code(traceback.format_exc())
                    st.warning("Try disabling charts or certain sections if the error persists.")
    
    def _extract_hull_metrics(self, vessel_data, hull_agent):
        try:
            # Calculate hull performance metrics from vessel data
            hull_metrics = {}
            
            # Filter for power loss data
            filtered_data_power = []
            for entry in vessel_data:
                if 'report_date' in entry and 'hull_roughness_power_loss' in entry:
                    if entry['hull_roughness_power_loss'] is not None:
                        filtered_data_power.append(entry)
            
            if filtered_data_power:
                # Sort data by date
                filtered_data_power.sort(key=lambda x: pd.to_datetime(x['report_date']))
                
                # Get the latest power loss value
                latest_entry = filtered_data_power[-1]
                power_loss = latest_entry.get('hull_roughness_power_loss', 0)
                
                # Use hull agent to get condition assessment
                condition, _ = hull_agent.get_hull_condition(power_loss)
                
                # Generate chart
                chart, latest_power_loss = hull_agent.create_performance_chart(
                    filtered_data_power, 
                    'hull_roughness_power_loss',
                    "Hull Roughness - Excess Power Trend",
                    "Excess Power (%)"
                )
                
                # Calculate fuel savings potential
                # Simple formula: 0.1 MT/D per 2% power loss above 15%
                if power_loss > 15:
                    fuel_savings = (power_loss - 15) * 0.05
                else:
                    fuel_savings = 0
                
                # Recommendation based on hull condition
                if power_loss < 15:
                    recommendation = "Hull and propeller performance is good. No action required."
                elif 15 <= power_loss < 25:
                    recommendation = "Consider hull cleaning at next convenient opportunity."
                else:
                    recommendation = "Hull cleaning recommended as soon as possible."
                
                hull_metrics = {
                    'power_loss': power_loss,
                    'condition': condition,
                    'fuel_savings': fuel_savings,
                    'recommendation': recommendation,
                    'hull_chart': chart
                }
            else:
                # Default values if no data available
                hull_metrics = {
                    'power_loss': 0,
                    'condition': 'Unknown',
                    'fuel_savings': 0,
                    'recommendation': 'Insufficient data to provide recommendation.',
                    'hull_chart': None
                }
            
            return hull_metrics
        except Exception as e:
            print(f"Error in _extract_hull_metrics: {str(e)}")
            # Return some basic defaults in case of error
            return {
                'power_loss': 0,
                'condition': 'Data Processing Error',
                'fuel_savings': 0,
                'recommendation': f'Error extracting data: {str(e)}',
                'hull_chart': None
            }
    
    def _extract_speed_metrics(self, vessel_data):
        try:
            # Extract speed-consumption data
            speed_metrics = {}
            
            # Filter data for speed-consumption analysis
            ballast_data = []
            laden_data = []
            
            for entry in vessel_data:
                if ('speed' in entry and entry['speed'] is not None and 
                    'normalised_consumption' in entry and entry['normalised_consumption'] is not None and
                    'loading_condition' in entry and entry['loading_condition'] is not None):
                    
                    if entry['loading_condition'].lower() == 'ballast':
                        ballast_data.append(entry)
                    elif entry['loading_condition'].lower() == 'laden':
                        laden_data.append(entry)
            
            # Ballast condition data
            ballast_speeds = []
            ballast_consumptions = []
            if ballast_data:
                ballast_speeds = [entry.get('speed', 0) for entry in ballast_data]
                ballast_consumptions = [entry.get('normalised_consumption', 0) for entry in ballast_data]
                
                # Create ballast chart
                ballast_chart = self._create_simple_scatter_chart(
                    ballast_speeds, 
                    ballast_consumptions,
                    "Speed vs. Consumption - Ballast"
                )
            else:
                ballast_chart = None
            
            # Laden condition data
            laden_speeds = []
            laden_consumptions = []
            if laden_data:
                laden_speeds = [entry.get('speed', 0) for entry in laden_data]
                laden_consumptions = [entry.get('normalised_consumption', 0) for entry in laden_data]
                
                # Create laden chart
                laden_chart = self._create_simple_scatter_chart(
                    laden_speeds, 
                    laden_consumptions,
                    "Speed vs. Consumption - Laden"
                )
            else:
                laden_chart = None
            
            speed_metrics = {
                'ballast_data': ballast_data,
                'ballast_speeds': ballast_speeds,
                'ballast_consumptions': ballast_consumptions,
                'ballast_chart': ballast_chart,
                'laden_data': laden_data,
                'laden_speeds': laden_speeds,
                'laden_consumptions': laden_consumptions,
                'laden_chart': laden_chart
            }
            
            return speed_metrics
        except Exception as e:
            print(f"Error in _extract_speed_metrics: {str(e)}")
            # Return basic defaults in case of error
            return {
                'ballast_data': [],
                'ballast_speeds': [],
                'ballast_consumptions': [],
                'ballast_chart': None,
                'laden_data': [],
                'laden_speeds': [],
                'laden_consumptions': [],
                'laden_chart': None
            }
    
    def _create_simple_scatter_chart(self, x_data, y_data, chart_title):
        if not x_data or not y_data:
            return None
        
        try:
            # Create the figure
            fig = go.Figure()
            
            # Add scatter plot
            fig.add_trace(go.Scatter(
                x=x_data,
                y=y_data,
                mode='markers',
                marker=dict(
                    size=10,
                    color='#48cae4',
                ),
                name='Operational Data'
            ))
            
            # Calculate polynomial fit if we have enough data points
            if len(x_data) > 2:
                # Sort data by x for proper curve fitting
                sorted_indices = sorted(range(len(x_data)), key=lambda i: x_data[i])
                x_sorted = [x_data[i] for i in sorted_indices]
                y_sorted = [y_data[i] for i in sorted_indices]
                
                # Fit a 2nd order polynomial
                try:
                    coeffs = np.polyfit(x_sorted, y_sorted, 2)
                    poly = np.poly1d(coeffs)
                    
                    # Generate points for the curve
                    x_smooth = np.linspace(min(x_sorted), max(x_sorted), 100)
                    y_smooth = poly(x_smooth)
                    
                    # Add the polynomial curve
                    fig.add_trace(go.Scatter(
                        x=x_smooth,
                        y=y_smooth,
                        mode='lines',
                        line=dict(color='#ff006e', width=3),
                        name='Polynomial Fit'
                    ))
                except Exception as e:
                    print(f"Could not generate trend line: {e}")
            
            # Add a unique ID to the figure's layout
            unique_id = str(uuid.uuid4())
            fig.layout.update(
                title=chart_title,
                xaxis_title="Speed (knots)",
                yaxis_title="Fuel Consumption (mt/day)",
                template="plotly_dark",
                height=400,
                width=500,
                margin=dict(l=40, r=40, t=40, b=40),
                uirevision=unique_id
            )
            
            return fig
        except Exception as e:
            print(f"Error creating chart: {str(e)}")
            return None
    
    def _save_chart_as_image(self, fig):
        if fig is None:
            return None
        
        try:
            # Create a deep copy of the figure
            fig_copy = copy.deepcopy(fig)
            
            # Create a temporary file to save the image
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp:
                # Save the figure as a PNG file
                fig_copy.write_image(temp.name, scale=2)
                return temp.name
        except Exception as e:
            print(f"Error saving chart image: {e}")
            return None
    
    def _generate_docx_report(self, vessel_name, report_date, analyst_name, hull_metrics, speed_metrics, options):
        try:
            # Create a new Document
            doc = Document()
            
            # Add title page
            doc.add_heading('Vessel Performance Summary', 0)
            
            # Add a table for vessel information
            table = doc.add_table(rows=3, cols=2)
            table.style = 'Table Grid'
            
            rows = table.rows
            rows[0].cells[0].text = 'Vessel Name:'
            rows[0].cells[1].text = vessel_name.upper()
            
            rows[1].cells[0].text = 'Prepared On:'
            rows[1].cells[1].text = str(report_date)
            
            rows[2].cells[0].text = 'Prepared By:'
            rows[2].cells[1].text = analyst_name
            
            doc.add_paragraph()
            doc.add_paragraph('_' * 80)
            doc.add_paragraph()
            
            # Create summary table at the top
            summary_table = doc.add_table(rows=3, cols=6)
            summary_table.style = 'Table Grid'
            
            # Set headers for the top summary table
            header_row = summary_table.rows[0].cells
            header_row[0].text = "Hull & Propeller"
            header_row[1].text = "Average"
            header_row[2].text = "Machinery"
            header_row[3].text = "Good"
            header_row[4].text = "Emissions"
            header_row[5].text = "CII Rating - A"
            
            # Set content for the middle row (placeholder for icons)
            middle_row = summary_table.rows[1].cells
            middle_row[0].text = "[Hull Icon]"
            middle_row[1].text = f"Potential Savings\n{hull_metrics['fuel_savings']:.1f} MT/D"
            middle_row[2].text = "[Machinery Icon]"
            middle_row[3].text = "Potential Savings\n-"
            middle_row[4].text = "[Emissions Icon]"
            middle_row[5].text = "Potential Improvement\n-"
            
            doc.add_paragraph()
            
            # Hull & Propeller Performance section (if enabled)
            if options['include_hull']:
                doc.add_heading('Hull & Propeller Performance', 1)
                
                # Create table for hull performance data
                hull_table = doc.add_table(rows=5, cols=2)
                hull_table.style = 'Table Grid'
                
                # Add data to table
                rows = hull_table.rows
                rows[0].cells[0].text = 'Additional Power Consumption'
                rows[0].cells[1].text = f"{hull_metrics['power_loss']:.1f} %"
                
                rows[1].cells[0].text = 'Potential Fuel Savings from Hull Cleaning & Propeller Polishing'
                rows[1].cells[1].text = f"{hull_metrics['fuel_savings']:.1f} MT/D"
                
                rows[2].cells[0].text = 'Hull & Propeller Condition'
                rows[2].cells[1].text = hull_metrics['condition']
                
                rows[3].cells[0].text = 'Recommendation'
                rows[3].cells[1].text = hull_metrics['recommendation']
                
                rows[4].cells[0].text = 'Forecasted Date of Hull cleaning'
                rows[4].cells[1].text = '-'
                
                doc.add_paragraph()
                
                # Hull Performance Analysis section
                doc.add_heading('Hull & Propeller Performance Analysis', 1)
                
                # Add hull performance chart if available and charts are enabled
                if options['include_charts'] and 'hull_chart' in hull_metrics and hull_metrics['hull_chart'] is not None:
                    try:
                        # Save chart as image
                        chart_path = self._save_chart_as_image(hull_metrics['hull_chart'])
                        
                        if chart_path:
                            # Add image to document
                            doc.add_picture(chart_path, width=Inches(6.0))
                            # Delete temporary file
                            try:
                                os.unlink(chart_path)
                            except:
                                pass  # Ignore deletion errors
                        else:
                            doc.add_paragraph("[Chart image could not be generated]")
                    except Exception as e:
                        print(f"Error adding hull chart to document: {e}")
                        doc.add_paragraph("[Hull performance chart will be inserted here]")
                else:
                    doc.add_paragraph("[Hull performance chart will be inserted here]")
                
                # Add notes about hull performance
                doc.add_heading('Notes:', 2)
                notes = doc.add_paragraph()
                notes.add_run('- ').bold = True
                notes.add_run(f"The vessel shows {hull_metrics['power_loss']:.1f}% additional power consumption, indicating a {hull_metrics['condition'].lower()} hull condition.")
                
                if hull_metrics['fuel_savings'] > 0:
                    notes2 = doc.add_paragraph()
                    notes2.add_run('- ').bold = True
                    notes2.add_run(f"Potential fuel savings of {hull_metrics['fuel_savings']:.1f} MT/D could be achieved through hull cleaning and propeller polishing.")
                
                doc.add_paragraph()
            
            # Speed Consumption section (if enabled)
            if options['include_speed']:
                doc.add_heading('Speed Consumption Profile', 1)
                
                # Only add charts if enabled
                if options['include_charts']:
                    # Create a 2-column table for charts
                    chart_table = doc.add_table(rows=1, cols=2)
                    chart_table.style = 'Table Grid'
                    
                    # Left column - Ballast condition
                    if 'ballast_chart' in speed_metrics and speed_metrics['ballast_chart'] is not None:
                        try:
                            # Save chart as image
                            ballast_chart_path = self._save_chart_as_image(speed_metrics['ballast_chart'])
                            
                            if ballast_chart_path:
                                # Add paragraph with heading for the chart
                                chart_table.cell(0, 0).paragraphs[0].add_run("Ballast Condition").bold = True
                                # Add image to table cell
                                chart_table.cell(0, 0).add_paragraph().add_run().add_picture(ballast_chart_path, width=Inches(3.0))
                                # Delete temporary file
                                try:
                                    os.unlink(ballast_chart_path)
                                except:
                                    pass  # Ignore deletion errors
                            else:
                                chart_table.cell(0, 0).add_paragraph("[Ballast chart could not be generated]")
                        except Exception as e:
                            print(f"Error adding ballast chart: {e}")
                            chart_table.cell(0, 0).add_paragraph("[Ballast condition chart]")
                    else:
                        chart_table.cell(0, 0).add_paragraph("[No ballast condition data available]")
                    
                    # Right column - Laden condition
                    if 'laden_chart' in speed_metrics and speed_metrics['laden_chart'] is not None:
                        try:
                            # Save chart as image
                            laden_chart_path = self._save_chart_as_image(speed_metrics['laden_chart'])
                            
                            if laden_chart_path:
                                # Add paragraph with heading for the chart
                                chart_table.cell(0, 1).paragraphs[0].add_run("Laden Condition").bold = True
                                # Add image to table cell
                                chart_table.cell(0, 1).add_paragraph().add_run().add_picture(laden_chart_path, width=Inches(3.0))
                                # Delete temporary file
                                try:
                                    os.unlink(laden_chart_path)
                                except:
                                    pass  # Ignore deletion errors
                            else:
                                chart_table.cell(0, 1).add_paragraph("[Laden chart could not be generated]")
                        except Exception as e:
                            print(f"Error adding laden chart: {e}")
                            chart_table.cell(0, 1).add_paragraph("[Laden condition chart]")
                    else:
                        chart_table.cell(0, 1).add_paragraph("[No laden condition data available]")
                else:
                    # Add placeholder text if charts are disabled
                    doc.add_paragraph("[Speed consumption charts will be inserted here]")
                
                doc.add_paragraph()
                
                doc.add_heading('Notes:', 2)
                speed_notes = doc.add_paragraph()
                speed_notes.add_run('- ').bold = True
                
                if 'ballast_data' in speed_metrics and speed_metrics['ballast_data']:
                    if speed_metrics['ballast_consumptions']:
                        ballast_avg = sum(speed_metrics['ballast_consumptions']) / len(speed_metrics['ballast_consumptions'])
                        speed_notes.add_run(f"In ballast condition, the vessel shows an average fuel consumption of {ballast_avg:.1f} MT/day.")
                    else:
                        speed_notes.add_run("Insufficient data to analyze ballast condition performance.")
                else:
                    speed_notes.add_run("Insufficient data to analyze ballast condition performance.")
                
                speed_notes2 = doc.add_paragraph()
                speed_notes2.add_run('- ').bold = True
                
                if 'laden_data' in speed_metrics and speed_metrics['laden_data']:
                    if speed_metrics['laden_consumptions']:
                        laden_avg = sum(speed_metrics['laden_consumptions']) / len(speed_metrics['laden_consumptions'])
                        speed_notes2.add_run(f"In laden condition, the vessel shows an average fuel consumption of {laden_avg:.1f} MT/day.")
                    else:
                        speed_notes2.add_run("Insufficient data to analyze laden condition performance.")
                else:
                    speed_notes2.add_run("Insufficient data to analyze laden condition performance.")
                
                doc.add_paragraph()
            
            # Emissions section (if enabled)
            if options['include_emissions']:
                doc.add_heading('Emissions Profile', 1)
                doc.add_paragraph("CII rating for 2024 of the vessel is 'A' (exclusions not included). CII rating for 2024 is provisional, as it is subject to further verification and adjustments based on exclusion data.")
                doc.add_paragraph("[Emissions chart will be inserted here]")
                doc.add_paragraph()
            
            # Machinery section (if enabled)
            if options['include_machinery']:
                doc.add_heading('Main Engine Performance', 1)
                machinery_table = doc.add_table(rows=3, cols=2)
                machinery_table.style = 'Table Grid'
                
                rows = machinery_table.rows
                rows[0].cells[0].text = 'Average ME SFOC'
                rows[0].cells[1].text = "Data not available"
                
                rows[1].cells[0].text = 'ME Recommendation'
                rows[1].cells[1].text = "ME Performance is within Acceptable Range"
                
                rows[2].cells[0].text = 'Potential Fuel Saving from ME'
                rows[2].cells[1].text = "-"
                
                doc.add_paragraph()
                
                doc.add_heading('Auxiliaries Performance', 2)
                aux_table = doc.add_table(rows=2, cols=2)
                aux_table.style = 'Table Grid'
                
                rows = aux_table.rows
                rows[0].cells[0].text = 'Excess Boiler Consumption (last 6 month)'
                rows[0].cells[1].text = "Data not available"
                
                rows[1].cells[0].text = 'Redundant AE Hours (last 6 month)'
                rows[1].cells[1].text = "-"
                
                doc.add_paragraph()
            
            # Appendix
            doc.add_heading('Appendix', 1)
            doc.add_heading('General Conditions', 2)
            
            appendix = doc.add_paragraph()
            appendix.add_run('- ').bold = True
            appendix.add_run("Analysis Period is Last Six Months or after the Last Event whichever is later")
            
            appendix2 = doc.add_paragraph()
            appendix2.add_run('- ').bold = True
            appendix2.add_run("Days with Good Weather (BF<=4) are considered for analysis.")
            
            appendix3 = doc.add_paragraph()
            appendix3.add_run('- ').bold = True
            appendix3.add_run("Days with Steaming hrs greater than 17 considered for analysis.")
            
            appendix4 = doc.add_paragraph()
            appendix4.add_run('- ').bold = True
            appendix4.add_run("Data is compared with Original Sea Trial")
            
            doc.add_heading('Hull Performance', 2)
            hull_app1 = doc.add_paragraph()
            hull_app1.add_run('- ').bold = True
            hull_app1.add_run("Excess Power < 15 % -- Rating Good")
            
            hull_app2 = doc.add_paragraph()
            hull_app2.add_run('- ').bold = True
            hull_app2.add_run("15 < Excess Power < 25 % -- Rating Average")
            
            hull_app3 = doc.add_paragraph()
            hull_app3.add_run('- ').bold = True
            hull_app3.add_run("Excess Power > 25 % -- Rating Poor")
            
            if options['include_machinery']:
                doc.add_heading('Machinery Performance', 2)
                mach_app1 = doc.add_paragraph()
                mach_app1.add_run('- ').bold = True
                mach_app1.add_run("SFOC(Grms/kW.hr) within +/- 10 from Shop test condition are considered as \"Good\"")
                
                mach_app2 = doc.add_paragraph()
                mach_app2.add_run('- ').bold = True
                mach_app2.add_run("SFOC(Grms/kW.hr) Greater than 10 and less than 20 are considered as \"Average\"")
                
                mach_app3 = doc.add_paragraph()
                mach_app3.add_run('- ').bold = True
                mach_app3.add_run("SFOC(Grms/kW.hr) Above 20 are considered as \"Poor\"")
            
            if options['include_speed']:
                doc.add_heading('Speed Consumption Performance', 2)
                speed_app = doc.add_paragraph()
                speed_app.add_run('- ').bold = True
                speed_app.add_run("Analysis carried out using regression techniques.")
            
            # Convert the document to bytes
            docx_file = BytesIO()
            doc.save(docx_file)
            docx_file.seek(0)
            
            return docx_file.getvalue()
        except Exception as e:
            # Log the error
            print(f"Error generating report: {str(e)}")
            print(traceback.format_exc())
            
            # Create a simple error report instead
            error_doc = Document()
            error_doc.add_heading('ERROR: Report Generation Failed', 0)
            error_doc.add_paragraph(f"Vessel Name: {vessel_name}")
            error_doc.add_paragraph(f"Date: {report_date}")
            
            error_doc.add_heading('Error Details', 1)
            error_doc.add_paragraph(f"An error occurred while generating the report: {str(e)}")
            
            error_doc.add_heading('Recommendations', 1)
            error_doc.add_paragraph("1. Try disabling charts in the report options.")
            error_doc.add_paragraph("2. Check that all required data is available for the vessel.")
            error_doc.add_paragraph("3. Ensure all dependencies are installed: python-docx, plotly, kaleido.")
            
            # Basic report content without charts
            error_doc.add_heading('Basic Report Content (No Charts)', 1)
            
            if options['include_hull']:
                error_doc.add_heading('Hull & Propeller Performance', 2)
                error_doc.add_paragraph(f"Additional Power Consumption: {hull_metrics['power_loss']:.1f} %")
                error_doc.add_paragraph(f"Potential Fuel Savings: {hull_metrics['fuel_savings']:.1f} MT/D")
                error_doc.add_paragraph(f"Hull & Propeller Condition: {hull_metrics['condition']}")
                error_doc.add_paragraph(f"Recommendation: {hull_metrics['recommendation']}")
            
            if options['include_speed']:
                error_doc.add_heading('Speed Consumption Summary', 2)
                if 'ballast_data' in speed_metrics and speed_metrics['ballast_data']:
                    if speed_metrics['ballast_consumptions']:
                        ballast_avg = sum(speed_metrics['ballast_consumptions']) / len(speed_metrics['ballast_consumptions'])
                        error_doc.add_paragraph(f"Average ballast consumption: {ballast_avg:.1f} MT/day")
                
                if 'laden_data' in speed_metrics and speed_metrics['laden_data']:
                    if speed_metrics['laden_consumptions']:
                        laden_avg = sum(speed_metrics['laden_consumptions']) / len(speed_metrics['laden_consumptions'])
                        error_doc.add_paragraph(f"Average laden consumption: {laden_avg:.1f} MT/day")
            
            # Convert the error document to bytes
            error_file = BytesIO()
            error_doc.save(error_file)
            error_file.seek(0)
            
            return error_file.getvalue()
