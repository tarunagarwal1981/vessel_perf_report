import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import base64
import tempfile
import os
import copy
import uuid
import traceback

class AdvancedReportGenerator:
    def __init__(self):
        # Define paths to icon images - update these paths to your actual icon locations
        self.icon_paths = {
            'hull_icon': 'icons/hull_icon.png',
            'machinery_icon': 'icons/machinery_icon.png',
            'emissions_icon': 'icons/emissions_icon.png'
        }
    
    def run(self, vessel_data, selected_vessel, hull_agent, speed_agent, cii_agent=None):
        st.header("Performance Report Generator")
        
        # Troubleshooting data
        self._troubleshoot_data(vessel_data)
        
        try:
            # Create tabs for report view and download
            report_tab1, report_tab2 = st.tabs(["Report Preview", "Download Report"])
            
            with report_tab1:
                # Display basic report preview
                st.subheader("Vessel Performance Summary")
                
                # Extract metrics from agents
                hull_metrics = self._get_hull_metrics_from_agent(vessel_data, hull_agent)
                speed_metrics = self._get_speed_metrics_from_agent(vessel_data, speed_agent)
                
                # Get CII metrics if CII agent is provided
                cii_metrics = None
                if cii_agent:
                    cii_metrics, _ = cii_agent.get_cii_data_for_report(vessel_data, selected_vessel)
                
                # Create preview layout similar to the document
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown("### Hull & Propeller")
                    st.markdown(f"**Condition:** {hull_metrics['condition']}")
                    st.markdown(f"**Power Loss:** {hull_metrics['power_loss']:.1f}%")
                    st.markdown(f"**Potential Savings:** {hull_metrics['fuel_savings']:.1f} MT/D")
                
                with col2:
                    st.markdown("### Speed & Consumption")
                    st.markdown(f"**Ballast:** {speed_metrics['ballast_avg']:.1f} MT/day")
                    st.markdown(f"**Laden:** {speed_metrics['laden_avg']:.1f} MT/day")
                
                with col3:
                    st.markdown("### Emissions")
                    if cii_metrics:
                        rating_color = self._get_cii_rating_color(cii_metrics['rating'])
                        st.markdown(f"**CII Rating:** <span style='color:{rating_color};font-weight:bold;'>{cii_metrics['rating']}</span>", unsafe_allow_html=True)
                        st.markdown(f"**AER:** {cii_metrics['attained_aer']:.2f} gCO₂/dwt-nm")
                    else:
                        st.markdown("**CII Rating:** No data")
                        st.markdown("**AER:** No data")
                
                # Display hull performance chart from hull agent
                st.subheader("Hull & Propeller Performance Analysis")
                # Get filtered data and generate chart using the existing agent
                hull_chart, latest_power_loss = self._get_hull_performance_chart_from_agent(vessel_data, hull_agent)
                if hull_chart:
                    st.plotly_chart(hull_chart, use_container_width=True, key=f"hull_preview_{str(uuid.uuid4())[:8]}")
                
                # Display speed-consumption charts from speed agent
                st.subheader("Speed Consumption Profile")
                col1, col2 = st.columns(2)
                
                ballast_chart, laden_chart = self._get_speed_consumption_charts_from_agent(vessel_data, speed_agent)
                
                with col1:
                    st.markdown("**Ballast Condition**")
                    if ballast_chart:
                        st.plotly_chart(ballast_chart, use_container_width=True, key=f"ballast_preview_{str(uuid.uuid4())[:8]}")
                    else:
                        st.info("No ballast data available")
                
                with col2:
                    st.markdown("**Laden Condition**")
                    if laden_chart:
                        st.plotly_chart(laden_chart, use_container_width=True, key=f"laden_preview_{str(uuid.uuid4())[:8]}")
                    else:
                        st.info("No laden data available")
                
                # Display CII chart if CII agent is provided
                if cii_agent:
                    st.subheader("Carbon Intensity Indicator (CII) Analysis")
                    _, cii_chart = cii_agent.get_cii_data_for_report(vessel_data, selected_vessel)
                    if cii_chart:
                        st.plotly_chart(cii_chart, use_container_width=True, key=f"cii_preview_{str(uuid.uuid4())[:8]}")
                    else:
                        st.info("No CII data available")
            
            with report_tab2:
                st.markdown("### Generate Report")
                
                # Report inputs
                col1, col2 = st.columns(2)
                
                with col1:
                    report_date = st.date_input("Report Date")
                    analyst_name = st.text_input("Analyst Name", "Marine Performance Team")
                
                with col2:
                    # Include sections options
                    st.subheader("Report Sections")
                    include_hull = st.checkbox("Include Hull Performance", value=True)
                    include_speed = st.checkbox("Include Speed Consumption", value=True)
                    include_emissions = st.checkbox("Include Emissions", value=True if cii_agent else False)
                    include_machinery = st.checkbox("Include Machinery (Placeholder)", value=False)
                
                # Report template options
                st.subheader("Report Template")
                template_option = st.radio(
                    "Select Template Style",
                    options=["Basic Template", "Advanced Template (With Icons)"],
                    index=1
                )
                
                # Generate report button
                if st.button("Generate Report"):
                    try:
                        with st.spinner("Generating report..."):
                            # Get metrics from agents for report
                            hull_metrics = self._get_hull_metrics_from_agent(vessel_data, hull_agent)
                            speed_metrics = self._get_speed_metrics_from_agent(vessel_data, speed_agent)
                            
                            # Generate the report document
                            docx_file = self._generate_formatted_report(
                                vessel_name=selected_vessel,
                                report_date=report_date,
                                analyst_name=analyst_name,
                                hull_metrics=hull_metrics,
                                speed_metrics=speed_metrics,
                                options={
                                    'include_hull': include_hull,
                                    'include_speed': include_speed,
                                    'include_emissions': include_emissions,
                                    'include_machinery': include_machinery,
                                    'template_option': template_option
                                },
                                vessel_data=vessel_data,
                                hull_agent=hull_agent,
                                speed_agent=speed_agent,
                                cii_agent=cii_agent
                            )
                            
                            # Create download button
                            st.download_button(
                                label="Download Report",
                                data=docx_file,
                                file_name=f"{selected_vessel}_Performance_Report_{report_date}.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key="download_report_button"
                            )
                            
                            st.success("Report generated successfully!")
                    except Exception as e:
                        st.error(f"Error generating report: {str(e)}")
                        st.code(traceback.format_exc())
        except Exception as e:
            st.error(f"Error in report generation UI: {str(e)}")
            st.code(traceback.format_exc())
    
    def _troubleshoot_data(self, data):
        """Analyze data for troubleshooting purposes"""
        # Check for speed consumption data
        has_speed = False
        has_consumption = False
        has_loading_condition = False
        has_hull_roughness = False
        ballast_count = 0
        laden_count = 0
        
        for entry in data:
            if 'speed' in entry and entry['speed'] is not None:
                has_speed = True
            
            if 'normalised_consumption' in entry and entry['normalised_consumption'] is not None:
                has_consumption = True
            
            if 'loading_condition' in entry and entry['loading_condition'] is not None:
                has_loading_condition = True
                
                # Count loading conditions
                if entry['loading_condition'].lower() == 'ballast':
                    ballast_count += 1
                elif entry['loading_condition'].lower() == 'laden':
                    laden_count += 1
            
            if 'hull_roughness_power_loss' in entry and entry['hull_roughness_power_loss'] is not None:
                has_hull_roughness = True
        
        with st.expander("Data Troubleshooting"):
            st.markdown(f"""
            **Data Summary:**
            - Total entries: {len(data)}
            - Has speed data: {has_speed}
            - Has consumption data: {has_consumption}
            - Has loading condition data: {has_loading_condition}
            - Has hull roughness data: {has_hull_roughness}
            - Ballast entries: {ballast_count}
            - Laden entries: {laden_count}
            """)
            
            # Show a sample of data
            if len(data) > 0:
                st.markdown("**Sample Data Entry:**")
                st.json(data[0])
    
    def _get_hull_metrics_from_agent(self, vessel_data, hull_agent):
        """Get hull metrics using the hull performance agent"""
        try:
            # Get hull chart to extract the latest power loss from the trend line
            _, latest_power_loss = self._get_hull_performance_chart_from_agent(vessel_data, hull_agent)
            
            if latest_power_loss is None:
                # Fallback if no trend line data is available
                # Filter data for power loss (same logic as in hull_agent)
                filtered_data = []
                for entry in vessel_data:
                    if 'report_date' in entry and 'hull_roughness_power_loss' in entry:
                        if entry['hull_roughness_power_loss'] is not None:
                            filtered_data.append(entry)
        
                if filtered_data:
                    # Sort by date
                    filtered_data.sort(key=lambda x: pd.to_datetime(x['report_date']))
        
                    # Get latest power loss
                    power_loss = filtered_data[-1].get('hull_roughness_power_loss', 0)
                else:
                    power_loss = 0
            else:
                # Use the value from the trend line
                power_loss = latest_power_loss
    
            # Get condition using the agent's method
            condition, _ = hull_agent.get_hull_condition(power_loss)
    
            # Calculate savings (formula can be adjusted)
            fuel_savings = (power_loss - 15) * 0.05 if power_loss > 15 else 0
    
            # Generate recommendation
            if power_loss < 15:
                recommendation = "Hull and propeller performance is good. No action required."
            elif 15 <= power_loss < 25:
                recommendation = "Consider hull cleaning at next convenient opportunity."
            else:
                recommendation = "Hull cleaning recommended as soon as possible."
    
            return {
                'condition': condition,
                'power_loss': power_loss,
                'fuel_savings': fuel_savings,
                'recommendation': recommendation
            }
        except Exception as e:
            # Log the error
            print(f"Error extracting hull metrics: {str(e)}")
            traceback.print_exc()
            return {
                'condition': "Error",
                'power_loss': 0,
                'fuel_savings': 0,
                'recommendation': f"Error analyzing hull performance: {str(e)}"
            }
    
    def _get_speed_metrics_from_agent(self, vessel_data, speed_agent):
        """Get speed metrics using the speed consumption agent's data"""
        try:
            # Filter data for speed-consumption
            ballast_consumptions = []
            laden_consumptions = []
            
            for entry in vessel_data:
                if 'loading_condition' in entry and 'normalised_consumption' in entry:
                    if entry['normalised_consumption'] is not None:
                        if entry.get('loading_condition', '').lower() == 'ballast':
                            ballast_consumptions.append(entry['normalised_consumption'])
                        elif entry.get('loading_condition', '').lower() == 'laden':
                            laden_consumptions.append(entry['normalised_consumption'])
            
            # Calculate averages
            ballast_avg = sum(ballast_consumptions) / len(ballast_consumptions) if ballast_consumptions else 0
            laden_avg = sum(laden_consumptions) / len(laden_consumptions) if laden_consumptions else 0
            
            return {
                'ballast_avg': ballast_avg,
                'laden_avg': laden_avg
            }
        except Exception as e:
            print(f"Error extracting speed metrics: {str(e)}")
            traceback.print_exc()
            return {
                'ballast_avg': 0,
                'laden_avg': 0
            }
    
    def _get_hull_performance_chart_from_agent(self, vessel_data, hull_agent):
        """Get hull performance chart from the hull agent with trend line value"""
        try:
            # Filter data for hull performance
            filtered_data = []
            for entry in vessel_data:
                if 'report_date' in entry and 'hull_roughness_power_loss' in entry:
                    if entry['hull_roughness_power_loss'] is not None:
                        filtered_data.append(entry)
            
            if filtered_data:
                # Use the agent's chart creation method
                chart, latest_value = hull_agent.create_performance_chart(
                    filtered_data,
                    'hull_roughness_power_loss',
                    "Hull Roughness - Excess Power Trend",
                    "Excess Power (%)"
                )
                
                # Optimize chart for full-width display
                if chart is not None:
                    # Adjust layout for better full-width display
                    chart.update_layout(
                        height=450,  # Slightly shorter height for better aspect ratio
                        margin=dict(l=50, r=50, t=60, b=50),  # More balanced margins
                        legend=dict(
                            orientation="h",
                            yanchor="bottom",
                            y=1.02,
                            xanchor="center",
                            x=0.5  # Center the legend
                        ),
                        # Increase font sizes for better readability
                        title=dict(font=dict(size=20)),
                        xaxis=dict(title=dict(font=dict(size=16))),
                        yaxis=dict(title=dict(font=dict(size=16)))
                    )
                
                return chart, latest_value
            
            return None, None
        except Exception as e:
            print(f"Error creating hull performance chart: {str(e)}")
            traceback.print_exc()
            return None, None
    
    def _get_speed_consumption_charts_from_agent(self, vessel_data, speed_agent):
        """Get speed consumption charts from the speed agent"""
        try:
            # Create ballast chart
            ballast_chart = None
            laden_chart = None
            
            # Check if we have valid data for charts
            filtered_data_consumption = []
            for entry in vessel_data:
                if ('speed' in entry and entry['speed'] is not None and 
                    'normalised_consumption' in entry and entry['normalised_consumption'] is not None and
                    'loading_condition' in entry and entry['loading_condition'] is not None):
                    filtered_data_consumption.append(entry)
            
            if filtered_data_consumption:
                # Generate ballast chart using speed_agent's method
                ballast_chart = speed_agent.create_speed_consumption_chart(
                    filtered_data_consumption,
                    "ballast",
                    "Speed vs. Consumption - Ballast Condition"
                )
                
                # Optimize ballast chart for side-by-side display
                if ballast_chart:
                    ballast_chart.update_layout(
                        height=400,  # Reduced height for side-by-side
                        margin=dict(l=40, r=40, t=50, b=50),
                        font=dict(size=12)  # Smaller font for side-by-side
                    )
                
                # Generate laden chart using speed_agent's method
                laden_chart = speed_agent.create_speed_consumption_chart(
                    filtered_data_consumption,
                    "laden",
                    "Speed vs. Consumption - Laden Condition"
                )
                
                # Optimize laden chart for side-by-side display
                if laden_chart:
                    laden_chart.update_layout(
                        height=400,  # Reduced height for side-by-side
                        margin=dict(l=40, r=40, t=50, b=50),
                        font=dict(size=12)  # Smaller font for side-by-side
                    )
            
            return ballast_chart, laden_chart
        except Exception as e:
            print(f"Error creating speed consumption charts: {str(e)}")
            traceback.print_exc()
            return None, None
    
    def _get_cii_rating_color(self, rating):
        """Get color code for CII rating"""
        rating_colors = {
            'A': '#4CAF50',  # Green
            'B': '#8BC34A',  # Light Green
            'C': '#FFC107',  # Amber
            'D': '#FF9800',  # Orange
            'E': '#F44336'   # Red
        }
        return rating_colors.get(rating, '#757575')  # Default to gray
    
    def _calculate_emissions_improvement(self, cii_data):
        """Calculate emissions improvement text for report"""
        if not cii_data:
            return "-"
            
        attained = cii_data['attained_aer']
        required = cii_data['required_cii']
        
        if attained < required:
            return f"The vessel's AER ({attained:.2f}) is below the required CII ({required:.2f}), demonstrating good performance."
        else:
            percent_improvement = ((attained - required) / attained) * 100
            return f"A {percent_improvement:.1f}% improvement in emissions is needed to reach the required CII."
    
    def _calculate_hull_impact(self, hull_metrics, cii_data):
        """Calculate hull impact on AER for report"""
        if not cii_data or not hull_metrics:
            return "-"
        
        power_loss = hull_metrics['power_loss']
        if power_loss <= 15:
            return "Hull condition is good and has minimal impact on emissions."
        else:
            # Estimate impact of hull condition on AER
            # Assuming a linear relationship for simplicity
            impact_percent = (power_loss - 15) * 0.5  # 0.5% impact per 1% excess power loss above 15%
            impact_aer = (cii_data['attained_aer'] * impact_percent) / 100
            
            return f"Hull condition contributes approximately {impact_aer:.2f} gCO₂/dwt-nm ({impact_percent:.1f}%) to the current AER."
    
    def _save_chart_as_image(self, fig):
        if fig is None:
            return None
        
        try:
            # Create a deep copy to avoid modifying the original
            fig_copy = copy.deepcopy(fig)
            
            # Update layout for better image output
            fig_copy.update_layout(
                template="plotly_white",  # Better for print
                height=600,
                width=1000,  # Increased width for better aspect ratio
                margin=dict(l=40, r=40, t=50, b=40),
                font=dict(color='black', size=14)
            )
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp:
                # Use higher scale for better resolution
                fig_copy.write_image(temp.name, scale=3)
                return temp.name
        except Exception as e:
            print(f"Error saving chart image: {str(e)}")
            traceback.print_exc()
            return None
    
    def _replace_text_in_document(self, doc, replacements):
        """Replace text in all parts of the document including headers/footers"""
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
        
        # Replace in headers and footers
        for section in doc.sections:
            # Headers
            for paragraph in section.header.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
            
            # Also check header tables
            for table in section.header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in replacements.items():
                                if key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(key, value)
            
            # Footers
            for paragraph in section.footer.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
            
            # Also check footer tables
            for table in section.footer.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in replacements.items():
                                if key in paragraph.text:
                                    paragraph.text = paragraph.text.replace(key, value)
    
    def _replace_chart_in_document(self, doc, placeholder, chart, width_inches=None):
        """
        Replace a placeholder in the document with a chart image
        
        Parameters:
        - doc: The Document object
        - placeholder: The text placeholder to replace
        - chart: The plotly chart object
        - width_inches: Optional width in inches (defaults to 6.0 for most charts or 7.5 for full width)
        """
        # Save chart as image
        chart_path = self._save_chart_as_image(chart)
        if not chart_path:
            return False
        
        # Default width if not specified
        if width_inches is None:
            width_inches = 6.0
        
        # Find and replace the placeholder with the image
        replaced = False
        try:
            # Try to replace in paragraphs
            for paragraph in doc.paragraphs:
                if placeholder in paragraph.text:
                    # Clear paragraph text
                    paragraph.text = ""
                    # Add picture
                    run = paragraph.add_run()
                    run.add_picture(chart_path, width=Inches(width_inches))
                    replaced = True
                    break
            
            # If not found in paragraphs, try tables
            if not replaced:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if placeholder in paragraph.text:
                                    # Clear paragraph text
                                    paragraph.text = ""
                                    # Add picture
                                    run = paragraph.add_run()
                                    # For tables, use a smaller width or adjust based on table cell
                                    cell_width = width_inches
                                    if cell_width > 3.5:  # Limit image width in table cells
                                        cell_width = 3.5
                                    run.add_picture(chart_path, width=Inches(cell_width))
                                    replaced = True
                                    break
                            if replaced:
                                break
                        if replaced:
                            break
                    if replaced:
                        break
            
            # Clean up
            try:
                os.unlink(chart_path)
            except:
                pass
            
            return replaced
        except Exception as e:
            print(f"Error replacing chart: {str(e)}")
            traceback.print_exc()
            return False
    
    def _generate_formatted_report(self, vessel_name, report_date, analyst_name, hull_metrics, 
                                 speed_metrics, options, vessel_data, hull_agent, speed_agent, cii_agent=None):
        try:
            # Add debug expander
            debug_expander = st.expander("Report Generation Debug")
            with debug_expander:
                st.write("Starting report generation...")
                st.write(f"Vessel name: {vessel_name}")
                st.write(f"Report date: {report_date}")
                st.write(f"CII agent provided: {cii_agent is not None}")
            
            # Set template path
            template_path = "templates/vessel_performance_template.docx"
            
            with debug_expander:
                st.write(f"Template path: {template_path}")
                st.write(f"Template exists: {os.path.exists(template_path)}")
                st.write(f"Current directory: {os.getcwd()}")
                
                try:
                    st.write(f"Files in current directory: {os.listdir()}")
                    if os.path.exists('templates'):
                        st.write(f"Files in templates directory: {os.listdir('templates')}")
                    else:
                        st.write("'templates' directory does not exist")
                except Exception as dir_error:
                    st.error(f"Error listing directory: {str(dir_error)}")
            
            # Try to open template or fallback
            try:
                if os.path.exists(template_path):
                    doc = Document(template_path)
                    with debug_expander:
                        st.write("✅ Template loaded successfully")
                else:
                    with debug_expander:
                        st.warning("Template not found, creating basic document instead")
                    # Create a basic document
                    doc = Document()
                    doc.add_heading('Vessel Performance Report', 0)
                    doc.add_paragraph(f"Vessel Name: {vessel_name}")
                    doc.add_paragraph(f"Date: {report_date}")
                    doc.add_paragraph(f"Hull Condition: {hull_metrics['condition']}")
                    doc.add_paragraph(f"Power Loss: {hull_metrics['power_loss']:.1f}%")
                    doc.add_paragraph(f"Potential Savings: {hull_metrics['fuel_savings']:.1f} MT/D")
            except Exception as e:
                with debug_expander:
                    st.error(f"Error opening template: {str(e)}")
                # Create a very basic document as fallback
                doc = Document()
                doc.add_heading('ERROR: Template Could Not Be Loaded', 0)
                doc.add_paragraph(f"Error: {str(e)}")
                doc.add_paragraph(f"Vessel Name: {vessel_name}")
                doc.add_paragraph(f"Date: {report_date}")
            
            # Try to replace placeholders
            try:
                with debug_expander:
                    st.write("Replacing text placeholders...")
                
                # Get CII metrics if agent is provided
                cii_data = None
                if cii_agent:
                    with debug_expander:
                        st.write("Fetching CII data for report...")
                    cii_metrics, _ = cii_agent.get_cii_data_for_report(vessel_data, vessel_name)
                    if cii_metrics:
                        cii_data = cii_metrics
                        with debug_expander:
                            st.write(f"✅ CII data retrieved: Rating = {cii_metrics['rating']}")
                
                replacements = {
                    '{{VESSEL_NAME}}': vessel_name.upper(),
                    '{{REPORT_DATE}}': report_date.strftime('%B %Y'),
                    '{{HULL_CONDITION}}': hull_metrics['condition'],
                    '{{HULL_SAVINGS}}': f"{hull_metrics['fuel_savings']:.1f} MT/D",
                    '{{POWER_CONSUMPTION}}': f"{hull_metrics['power_loss']:.1f} %",
                    '{{HULL_RECOMMENDATION}}': hull_metrics['recommendation'],
                    '{{HULL_CONDITION_DETAIL}}': hull_metrics['condition'],
                    '{{FUEL_SAVINGS}}': f"{hull_metrics['fuel_savings']:.1f} MT/D",
                    '{{HULL_CLEANING_DATE}}': "-",
                    '{{BALLAST_AVG}}': f"{speed_metrics['ballast_avg']:.1f}",
                    '{{LADEN_AVG}}': f"{speed_metrics['laden_avg']:.1f}",
                }
                
                # Add CII data to replacements if available
                if cii_data:
                    replacements.update({
                        '{{CII_RATING}}': cii_data['rating'],
                        '{{CII_RATING_DETAIL}}': f"{cii_data['rating']} (AER: {cii_data['attained_aer']:.2f})",
                        '{{AER_VALUE}}': f"{cii_data['attained_aer']:.2f}",
                        '{{REQUIRED_CII}}': f"{cii_data['required_cii']:.2f}",
                        '{{EMISSIONS_IMPROVEMENT}}': self._calculate_emissions_improvement(cii_data),
                        '{{HULL_IMPACT_ON_AER}}': self._calculate_hull_impact(hull_metrics, cii_data)
                    })
                else:
                    # Default values if CII data is not available
                    replacements.update({
                        '{{CII_RATING}}': "N/A",
                        '{{CII_RATING_DETAIL}}': "N/A",
                        '{{AER_VALUE}}': "N/A",
                        '{{REQUIRED_CII}}': "N/A",
                        '{{EMISSIONS_IMPROVEMENT}}': "-",
                        '{{HULL_IMPACT_ON_AER}}': "-"
                    })
                
                # Add remaining placeholders
                # Add remaining placeholders
                replacements.update({
                    '{{ME_SFOC}}': "167.12 g/KWhr at 81% Load (Placeholder)",
                    '{{ME_RECOMMENDATION}}': "ME Performance is within Acceptable Range",
                    '{{ME_FUEL_SAVING}}': "-",
                    '{{BOILER_CONSUMPTION}}': "16.7 MT",
                    '{{REDUNDANT_AE_HOURS}}': "-",
                    '{{HULL_NOTES}}': "The vessel tends to operate with a gradual change in added resistance over time.",
                    '{{SPEED_CONSUMPTION_NOTES}}': "In laden and ballast conditions, consumption remains relatively consistent."
                })
                
                self._replace_text_in_document(doc, replacements)
                with debug_expander:
                    st.write("✅ Text placeholders replaced")
            except Exception as e:
                with debug_expander:
                    st.error(f"Error replacing text: {str(e)}")
                    st.code(traceback.format_exc())
            
            # Try to replace charts
            try:
                with debug_expander:
                    st.write("Generating and replacing charts...")
                
                # Generate hull performance chart using the agent
                hull_chart, latest_power_loss = self._get_hull_performance_chart_from_agent(vessel_data, hull_agent)
                if hull_chart:
                    # Use 7.5 inches for full page width (typical A4 page width is about 8.3 inches, leaving margins)
                    success = self._replace_chart_in_document(doc, '{{HULL_PERFORMANCE_CHART}}', hull_chart, width_inches=7.5)
                    with debug_expander:
                        if success:
                            st.write("✅ Hull performance chart replaced (full width)")
                        else:
                            st.warning("⚠️ Hull performance chart placeholder not found")
                else:
                    with debug_expander:
                        st.warning("⚠️ No hull performance chart data available")
                
                # Generate and replace speed consumption charts using the agent
                ballast_chart, laden_chart = self._get_speed_consumption_charts_from_agent(vessel_data, speed_agent)
                
                # Debug speed charts
                with debug_expander:
                    st.write(f"Ballast chart available: {ballast_chart is not None}")
                    st.write(f"Laden chart available: {laden_chart is not None}")
                
                if ballast_chart:
                    # Use 3.5 inches width for side-by-side charts (7 inches total width)
                    success = self._replace_chart_in_document(doc, '{{BALLAST_CHART}}', ballast_chart, width_inches=3.5)
                    with debug_expander:
                        if success:
                            st.write("✅ Ballast chart replaced (side-by-side)")
                        else:
                            st.warning("⚠️ Ballast chart placeholder not found or could not be replaced")
                else:
                    with debug_expander:
                        st.warning("⚠️ No ballast chart data available")
                
                if laden_chart:
                    # Use 3.5 inches width for side-by-side charts (7 inches total width)
                    success = self._replace_chart_in_document(doc, '{{LADEN_CHART}}', laden_chart, width_inches=3.5)
                    with debug_expander:
                        if success:
                            st.write("✅ Laden chart replaced (side-by-side)")
                        else:
                            st.warning("⚠️ Laden chart placeholder not found or could not be replaced")
                else:
                    with debug_expander:
                        st.warning("⚠️ No laden chart data available")
                
                # Replace CII chart if CII agent is available
                if cii_agent and options.get('include_emissions', False):
                    with debug_expander:
                        st.write("Generating CII chart for report...")
                        
                    _, cii_chart = cii_agent.get_cii_data_for_report(vessel_data, vessel_name)
                    
                    if cii_chart:
                        # Use 7.5 inches for full page width
                        success = self._replace_chart_in_document(doc, '{{CII_CHART}}', cii_chart, width_inches=7.5)
                        with debug_expander:
                            if success:
                                st.write("✅ CII chart replaced (full width)")
                            else:
                                st.warning("⚠️ CII chart placeholder not found")
                    else:
                        with debug_expander:
                            st.warning("⚠️ No CII chart data available")
                
            except Exception as e:
                with debug_expander:
                    st.error(f"Error replacing charts: {str(e)}")
                    st.code(traceback.format_exc())
            
            # Save to memory
            with debug_expander:
                st.write("Saving document...")
            
            docx_file = BytesIO()
            doc.save(docx_file)
            docx_file.seek(0)
            
            with debug_expander:
                st.write("✅ Document saved successfully")
            
            return docx_file.getvalue()
        except Exception as e:
            # Log the error
            st.error(f"Error generating report: {str(e)}")
            st.code(traceback.format_exc())
            
            # Create a simple error report
            error_doc = Document()
            error_doc.add_heading('ERROR: Report Generation Failed', 0)
            error_doc.add_paragraph(f"Vessel Name: {vessel_name}")
            error_doc.add_paragraph(f"Date: {report_date}")
            
            error_doc.add_heading('Error Details', 1)
            error_doc.add_paragraph(f"An error occurred during report generation: {str(e)}")
            
            error_doc.add_heading('Basic Report Content (No Formatting)', 1)
            
            # Add basic content without fancy formatting
            error_doc.add_paragraph(f"Hull Condition: {hull_metrics['condition']}")
            error_doc.add_paragraph(f"Power Loss: {hull_metrics['power_loss']:.1f}%")
            error_doc.add_paragraph(f"Potential Savings: {hull_metrics['fuel_savings']:.1f} MT/D")
            
            if speed_metrics['ballast_avg'] > 0:
                error_doc.add_paragraph(f"Ballast Consumption: {speed_metrics['ballast_avg']:.1f} MT/day")
            
            if speed_metrics['laden_avg'] > 0:
                error_doc.add_paragraph(f"Laden Consumption: {speed_metrics['laden_avg']:.1f} MT/day")
            
            # Add CII information if available
            if cii_agent:
                try:
                    cii_metrics, _ = cii_agent.get_cii_data_for_report(vessel_data, vessel_name)
                    if cii_metrics:
                        error_doc.add_paragraph(f"CII Rating: {cii_metrics['rating']}")
                        error_doc.add_paragraph(f"Attained AER: {cii_metrics['attained_aer']:.2f} gCO₂/dwt-nm")
                        error_doc.add_paragraph(f"Required CII: {cii_metrics['required_cii']:.2f} gCO₂/dwt-nm")
                except Exception as cii_error:
                    error_doc.add_paragraph(f"Error getting CII data: {str(cii_error)}")
            
            # Save the error document
            error_file = BytesIO()
            error_doc.save(error_file)
            error_file.seek(0)
            
            return error_file.getvalue()
